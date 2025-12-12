from docx import Document
import re
import argparse
from tqdm import tqdm


# --------- HELPER FUNCTIONS ---------
def clean_text(text):
    """
    Clean up extracted text for RAG use:
    - Remove multiple empty lines
    - Remove page numbers, repeated titles, and macros
    - Strip leading/trailing whitespace
    """
    # Remove lines that are just numbers (page numbers)
    text = re.sub(r'^\s*\d+\s*$', '', text, flags=re.MULTILINE)
    
    # Remove common Confluence DOC artifacts like [TOC], [Macro], etc.
    text = re.sub(r'\[.*?\]', '', text)
    
    # Remove extra empty lines
    text = re.sub(r'\n\s*\n+', '\n\n', text)
    
    # Replace weird whitespace characters
    text = text.replace('\xa0', ' ').replace('\u200b', '')  # non-breaking & zero-width spaces
    
    # Strip leading/trailing whitespace
    return text.strip()

def extract_doc_text(doc_path):
    doc = Document(doc_path)
    full_text = []

    # ----- Remove headers & footers -----
    for section in doc.sections:
        # Headers
        for para in section.header.paragraphs:
            text = para.text.strip()
            if text:
                full_text.append(f"[HEADER REMOVED: {text}]")  # Optional: log removed header
        # Footers
        for para in section.footer.paragraphs:
            text = para.text.strip()
            if text:
                full_text.append(f"[FOOTER REMOVED: {text}]")  # Optional: log removed footer

    # ----- Extract main content with progress bar -----
    print("Extracting paragraphs...")
    for para in tqdm(doc.paragraphs, desc="Paragraphs", unit="para"):
        text = para.text.strip()
        if text:
            # Optional: preserve heading levels
            if para.style.name.startswith('Heading'):
                level = para.style.name.replace('Heading ', '')
                text = f"\n=== Heading {level}: {text} ===\n"
            full_text.append(text)

    # ----- Extract tables with progress bar -----
    print("Extracting tables...")
    for table in tqdm(doc.tables, desc="Tables", unit="table"):
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                full_text.append(row_text)

    combined_text = "\n".join(full_text)
    return clean_text(combined_text)

# --------- MAIN ---------
def main():
    parser = argparse.ArgumentParser(description="Parse Confluence DOCx page to txt for RAG extraction")
    parser.add_argument("--input", required=True, help="Input DOCx file")
    parser.add_argument("--output", help="Output TXT file", default="output.txt")
    args = parser.parse_args()
    
    text = extract_doc_text(args.input)
    with open(args.output, "w", encoding="utf-8") as f:
        f.write(text)
        
    print(f"RAG-ready text saved to {args.output}")

if __name__ == "__main__":
    main()
